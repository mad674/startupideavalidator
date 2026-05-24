from docx import Document


class DOCXExtractor:

    def extract(self, path: str):

        doc = Document(path)

        text = []

        for para in doc.paragraphs:

            if para.text.strip():

                text.append(para.text)

        return {
            "text": "\n".join(text),
            "metadata": {
                "source_type": "docx",
                "paragraphs": len(doc.paragraphs)
            }
        }